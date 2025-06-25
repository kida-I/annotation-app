import streamlit as st
import pandas as pd
import uuid
import json
from io import StringIO, BytesIO
import docx  # python-docx
import openpyxl  # openpyxl
import re

# Page configuration
st.set_page_config(
    page_title="Interactive Text Annotation Tool",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .text-passage {
        background-color: #f0f2f6;
        padding: 15px;
        border-radius: 3px;
        border-left: 2px solid #1f77b4;
        margin-bottom: 5px;
    }
    .annotation-section {
        background-color: #ffffff;
        padding: 5px;
        border-radius: 3px;
        border: 1px solid #e0e0e0;
        margin-bottom: 5px;
    }
    .section-header {
        color: #1f77b4;
        font-weight: bold;
        margin-bottom: 5px;
    }
    .stDataEditor {
        border: 1px solid #e0e0e0;
        border-radius: 3px;
    }
    /* Buttons kleiner machen */
    button[kind="secondary"], button[kind="primary"], .stButton>button {
        font-size: 15px !important;
        padding: 2px 8px !important;
        height: 8px !important;
        min-width: 0 !important;
        border-radius: 3px !important;
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state with default data if not exists"""
    
    if 'text_passages' not in st.session_state:
        st.session_state.text_passages = [
            "Dies ist der erste Beispielabsatz. Er enthält grundlegende Informationen.",
            "Dieser zweite Absatz bietet ergänzende Details und Kontext zum Thema."
        ]
    
    if 'annotations' not in st.session_state:
        # Initialize annotations for each passage
        st.session_state.annotations = {
            0: pd.DataFrame({
                'Citation': [
                    'Beispielabsatz',
                    'grundlegende Informationen'
                ],
                'Metadata': [
                    'Author: Mustermann, Year: 2023, Page: 45',
                    'Author: Schmidt, Year: 2022, Chapter: 3'
                ],
                'Annotation_Text': [
                    'Diese Passage führt in das Thema ein und stellt die Grundlagen dar.',
                    'Der Begriff "grundlegende Informationen" bezieht sich auf Basisdaten.'
                ]
            }),
            1: pd.DataFrame({
                'Citation': [
                    'ergänzende Details'
                ],
                'Metadata': [
                    'Author: Weber, Year: 2024, Volume: 2, Page: 78'
                ],
                'Annotation_Text': [
                    'Dieser Absatz erweitert die Grundlagen um wichtige Kontextinformationen.'
                ]
            })
        }

def check_citation_in_text(citation, text):
    """Check if citation is found in the text"""
    if not citation or not text:
        return False
    return citation.lower() in text.lower()

def add_found_in_text_column(annotations_df, passage_text):
    """Add 'Found in Text' column to annotations DataFrame"""
    annotations_df = annotations_df.copy()
    annotations_df['Found in Text'] = annotations_df['Citation'].apply(
        lambda citation: "✅ Yes" if check_citation_in_text(citation, passage_text) else "❌ No"
    )
    return annotations_df

def add_annotation_row(passage_index):
    """Add a new empty row to the annotations for a specific passage"""
    new_row = pd.DataFrame({
        'Citation': [''],
        'Metadata': [''],
        'Annotation_Text': ['']
    })
    st.session_state.annotations[passage_index] = pd.concat(
        [st.session_state.annotations[passage_index], new_row], 
        ignore_index=True
    )

def remove_annotation_row(passage_index, row_index):
    """Remove a specific row from annotations"""
    if len(st.session_state.annotations[passage_index]) > row_index:
        st.session_state.annotations[passage_index] = st.session_state.annotations[passage_index].drop(
            index=row_index
        ).reset_index(drop=True)

def add_text_passage():
    """Add a new text passage"""
    st.session_state.text_passages.append("Neuer Textabsatz...")
    new_index = len(st.session_state.text_passages) - 1
    st.session_state.annotations[new_index] = pd.DataFrame({
        'Citation': [''],
        'Metadata': [''],
        'Annotation_Text': ['']
    })

def remove_text_passage(index):
    """Remove a text passage and its annotations"""
    if len(st.session_state.text_passages) > 1:
        st.session_state.text_passages.pop(index)
        # Remove annotations for this passage
        if index in st.session_state.annotations:
            del st.session_state.annotations[index]
        # Reindex remaining annotations
        new_annotations = {}
        for i, passage in enumerate(st.session_state.text_passages):
            old_index = i if i < index else i + 1
            if old_index in st.session_state.annotations:
                new_annotations[i] = st.session_state.annotations[old_index]
            else:
                new_annotations[i] = pd.DataFrame({
                    'Citation': [''],
                    'Metadata': [''],
                    'Annotation_Text': ['']
                })
        st.session_state.annotations = new_annotations

def export_data():
    """Export all data as a structured dictionary"""
    export_data = {
        'text_passages': st.session_state.text_passages,
        'annotations': {}
    }
    
    for passage_index, annotations_df in st.session_state.annotations.items():
        # Only export the core columns, not the computed 'Found in Text' column
        core_columns = ['Citation', 'Metadata', 'Annotation_Text']
        export_df = annotations_df[core_columns] if all(col in annotations_df.columns for col in core_columns) else annotations_df
        export_data['annotations'][passage_index] = export_df.to_dict('records')
    
    return export_data

def save_data():
    """Exportiere die aktuellen Daten als JSON-String"""
    data = export_data()
    return json.dumps(data, ensure_ascii=False, indent=2)

def load_data(json_str):
    """Lade Daten aus einem JSON-String in den Session State"""
    try:
        data = json.loads(json_str)
        st.session_state.text_passages = data.get('text_passages', [])
        # Annotations wieder als DataFrames speichern
        st.session_state.annotations = {}
        for idx, records in data.get('annotations', {}).items():
            try:
                idx_int = int(idx)
            except Exception:
                idx_int = idx
            st.session_state.annotations[idx_int] = pd.DataFrame(records)
        return True
    except json.JSONDecodeError as e:
        st.error(f"Fehler beim Laden der JSON-Datei: {e}")
        return False

def highlight_citations(text, citations):
    """Highlight all citation occurrences in text"""
    highlighted = text
    # Sort by length so longer citations are replaced first
    for citation in sorted(citations, key=len, reverse=True):
        if citation and citation in highlighted:
            highlighted = highlighted.replace(
                citation,
                f'<mark style="background-color: #b9f6ca; padding: 2px 4px; border-radius: 3px;">{citation}</mark>'
            )
    return highlighted

def highlight_citations_for_export(text, citations):
    """Fügt <span style="color:green">...</span> für alle Citation-Vorkommen im Text ein (für Word-Export)"""
    highlighted = text
    for citation in sorted(citations, key=len, reverse=True):
        if citation and citation in highlighted:
            highlighted = highlighted.replace(
                citation,
                f'<span style="color:green">{citation}</span>'
            )
    return highlighted

def import_word(file):
    """Importiere Text und Annotationen aus einem Word-Dokument (.docx)"""
    doc = docx.Document(file)
    passages = []
    annotations = {}
    current_passage = None
    current_annots = []
    # Erkenne Absätze und Annotationen
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Prüfe, ob es ein neuer Passage-Absatz ist (keine Einrückung, kein Bullet)
        if not para.style.name.startswith('List'):
            # Wenn vorher Passage und Annotationen gesammelt wurden, speichern
            if current_passage is not None:
                if current_annots:
                    annotations[len(passages)-1] = pd.DataFrame(current_annots)
                else:
                    annotations[len(passages)-1] = pd.DataFrame({
                        'Citation': [''],
                        'Metadata': [''],
                        'Annotation_Text': ['']
                    })
                current_annots = []
            passages.append(text)
            current_passage = text
        else:
            # Annotationen erkennen (Bullet-Listen)
            if text.startswith("Citation:"):
                citation = text[len("Citation:"):].strip()
                current_annots.append({'Citation': citation, 'Metadata': '', 'Annotation_Text': ''})
            elif text.startswith("Metadata:"):
                metadata = text[len("Metadata:"):].strip()
                if current_annots:
                    current_annots[-1]['Metadata'] = metadata
            elif text.startswith("Annotation:"):
                annotation = text[len("Annotation:"):].strip()
                if current_annots:
                    current_annots[-1]['Annotation_Text'] = annotation
    # Letzte Passage speichern
    if current_passage is not None:
        if current_annots:
            annotations[len(passages)-1] = pd.DataFrame(current_annots)
        else:
            annotations[len(passages)-1] = pd.DataFrame({
                'Citation': [''],
                'Metadata': [''],
                'Annotation_Text': ['']
            })
    # Fallback falls keine Annotationen erkannt wurden
    if not annotations:
        # ...alte Fußnoten-Logik...
        metadata = []
        for para in doc.paragraphs:
            if para.text.strip():
                notes = re.findall(r'\[(\d+)\]', para.text)
                for note in notes:
                    for p in doc.paragraphs:
                        if p.text.strip().startswith(f"{note} "):
                            metadata.append(p.text.strip())
        if not metadata:
            metadata = ["" for _ in passages]
        for i, passage in enumerate(passages):
            annotations[i] = pd.DataFrame({
                'Citation': [''],
                'Metadata': [metadata[i] if i < len(metadata) else ""],
                'Annotation_Text': ['']
            })
    return passages, annotations

def import_excel(file):
    """Importiere Text und Annotationen aus einer Excel-Datei"""
    df = pd.read_excel(file)
    # Erwartet: Spalten 'Text', 'Citation', 'Metadata', 'Annotation_Text'
    passages = []
    annotations = {}
    if 'Text' in df.columns:
        grouped = df.groupby('Text')
        for i, (text, group) in enumerate(grouped):
            passages.append(text)
            annotations[i] = group[['Citation', 'Metadata', 'Annotation_Text']].reset_index(drop=True)
    else:
        # Fallback: Jede Zeile ist ein Passage
        for i, row in df.iterrows():
            passages.append(str(row.get('Text', '')))
            annotations[i] = pd.DataFrame({
                'Citation': [row.get('Citation', '')],
                'Metadata': [row.get('Metadata', '')],
                'Annotation_Text': [row.get('Annotation_Text', '')]
            })
    return passages, annotations

def export_word(passages, annotations):
    """Exportiere Text und Annotationen in ein Word-Dokument (.docx) mit grüner Markierung für Citations"""
    doc = docx.Document()
    for i, passage in enumerate(passages):
        annots = annotations.get(i, pd.DataFrame())
        # Hole alle Citations für diese Passage
        citations = []
        if 'Citation' in annots.columns:
            citations = [c for c in annots['Citation'] if c and str(c).strip()]
        # Markiere Citations im Passage-Text
        highlighted_passage = highlight_citations_for_export(passage, citations)
        # Füge als HTML ein (python-docx unterstützt kein echtes HTML, daher Workaround: ersetze <span style="color:green">...</span> durch Run mit Farbe)
        p = doc.add_paragraph()
        last_idx = 0
        import re
        # Finde alle markierten Bereiche
        for match in re.finditer(r'<span style="color:green">(.*?)</span>', highlighted_passage):
            # Text vor dem markierten Bereich
            if match.start() > last_idx:
                p.add_run(highlighted_passage[last_idx:match.start()])
            # Markierter Bereich
            run = p.add_run(match.group(1))
            run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)  # grün
            last_idx = match.end()
        # Restlicher Text
        if last_idx < len(highlighted_passage):
            p.add_run(highlighted_passage[last_idx:])
        # Annotationen wie gehabt
        for idx, row in annots.iterrows():
            doc.add_paragraph(f"  Citation: {row.get('Citation','')}", style='List Bullet')
            doc.add_paragraph(f"  Metadata: {row.get('Metadata','')}", style='List Bullet')
            doc.add_paragraph(f"  Annotation: {row.get('Annotation_Text','')}", style='List Bullet')
        doc.add_paragraph("")  # Leerzeile
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_excel(passages, annotations):
    """Exportiere Text und Annotationen in eine Excel-Datei (.xlsx) mit fetter Markierung für Citations"""
    rows = []
    for i, passage in enumerate(passages):
        annots = annotations.get(i, pd.DataFrame())
        # Hole alle Citations für diese Passage
        citations = []
        if 'Citation' in annots.columns:
            citations = [c for c in annots['Citation'] if c and str(c).strip()]
        # Markiere Citations im Passage-Text mit **...** (für nachträgliche Formatierung)
        highlighted_passage = passage
        for citation in sorted(citations, key=len, reverse=True):
            if citation and citation in highlighted_passage:
                highlighted_passage = highlighted_passage.replace(
                    citation,
                    f'**{citation}**'
                )
        if annots.empty:
            rows.append({'Text': highlighted_passage, 'Citation': '', 'Metadata': '', 'Annotation_Text': ''})
        else:
            for idx, row in annots.iterrows():
                rows.append({
                    'Text': highlighted_passage,
                    'Citation': row.get('Citation', ''),
                    'Metadata': row.get('Metadata', ''),
                    'Annotation_Text': row.get('Annotation_Text', '')
                })
    df = pd.DataFrame(rows)
    output = BytesIO()
    # Schreibe DataFrame in Excel
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
        # Fette die Citation-Vorkommen in der Text-Spalte
        ws = writer.sheets['Sheet1']
        for row_idx, text in enumerate(df['Text'], start=2):  # Header ist Zeile 1
            if '**' in str(text):
                # Splitte und setze Runs mit Fett
                from openpyxl.styles import Font
                cell = ws.cell(row=row_idx, column=1)
                parts = re.split(r'(\*\*.*?\*\*)', str(text))
                cell.value = None  # Leere Zelle, wir setzen gleich RichText
                rich_text = []
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = openpyxl.cell.rich_text.Text(part[2:-2], font=Font(bold=True))
                    else:
                        run = openpyxl.cell.rich_text.Text(part)
                    rich_text.append(run)
                # openpyxl unterstützt ab v3.1 RichText (ansonsten: alles fett setzen)
                try:
                    cell.rich_text = rich_text
                except Exception:
                    # Fallback: alles fett, wenn RichText nicht unterstützt
                    cell.value = re.sub(r'\*\*(.*?)\*\*', r'\1', str(text))
        writer.save()
    output.seek(0)
    return output

def main():
    """Main application function"""
    
    # Initialize session state
    initialize_session_state()
    
    # Application header
    st.title("📝 Interactive Text Annotation Tool")
    st.markdown("---")
    
    # Control buttons in sidebar
    with st.sidebar:
        st.header("🛠️ Controls")
        
        # --- Speichern ---
        st.markdown("### 💾 Speichern")
        if st.button("💾 Daten als JSON speichern", use_container_width=True, key="save_json_btn"):
            json_data = save_data()
            st.download_button(
                label="⬇️ Download JSON",
                data=json_data,
                file_name="annotations.json",
                mime="application/json",
                use_container_width=True,
                key="download_json_btn"
            )
        # Export als Word
        if st.button("⬇️ Export als Word (.docx)", use_container_width=True, key="export_word_btn"):
            word_file = export_word(st.session_state.text_passages, st.session_state.annotations)
            st.download_button(
                label="Download Word",
                data=word_file,
                file_name="annotations.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_word_btn"
            )
        # Export als Excel
        if st.button("⬇️ Export als Excel (.xlsx)", use_container_width=True, key="export_excel_btn"):
            excel_file = export_excel(st.session_state.text_passages, st.session_state.annotations)
            st.download_button(
                label="Download Excel",
                data=excel_file,
                file_name="annotations.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="download_excel_btn"
            )

        # --- Laden ---
        st.markdown("### 📂 Laden")
        uploaded_file = st.file_uploader(
            "Datei laden (JSON, Word, Excel)",
            type=["json", "docx", "xlsx"],
            accept_multiple_files=False,
            key="file_uploader"
        )
        if uploaded_file is not None and uploaded_file.size > 0:
            if uploaded_file.name.endswith(".json"):
                json_str = uploaded_file.getvalue().decode("utf-8")
                if st.button("🔄 Daten laden (JSON)", use_container_width=True, key="load_json_btn"):
                    if load_data(json_str):
                        st.success("Daten erfolgreich geladen! Die Seite wird neu geladen.")
                        st.rerun()
            elif uploaded_file.name.endswith(".docx"):
                if st.button("🔄 Daten laden (Word)", use_container_width=True, key="load_word_btn"):
                    passages, annotations = import_word(uploaded_file)
                    st.session_state.text_passages = passages
                    st.session_state.annotations = annotations
                    st.success("Word-Dokument erfolgreich importiert! Die Seite wird neu geladen.")
                    st.rerun()
            elif uploaded_file.name.endswith(".xlsx"):
                if st.button("🔄 Daten laden (Excel)", use_container_width=True, key="load_excel_btn"):
                    passages, annotations = import_excel(uploaded_file)
                    st.session_state.text_passages = passages
                    st.session_state.annotations = annotations
                    st.success("Excel-Datei erfolgreich importiert! Die Seite wird neu geladen.")
                    st.rerun()
        if st.button("➕ Add Text Passage", use_container_width=True):
            add_text_passage()
            st.rerun()
        
        st.markdown("### 📊 Export Data")
        if st.button("📥 Show Export Data", use_container_width=True):
            st.json(export_data())
        
        st.markdown("### ℹ️ Instructions")
        st.markdown("""
        **How to use:**
        1. Edit text passages directly in the text areas
        2. Modify annotations in the tables
        3. Use ➕/➖ buttons to add/remove annotations
        4. Add new text passages with the sidebar button
        5. Speichern/Laden über die Sidebar
        6. **NEU:** Die Spalte "Found in Text" zeigt automatisch an, ob Citations im Text gefunden werden
        """)
    
    # Main content area
    for passage_index, passage_text in enumerate(st.session_state.text_passages):
        
        # Create columns for text and annotations
        col1, col2 = st.columns([1, 1], gap="large")
        
        with col1:
            # Text passage section
            # Entferne die Überschrift mit Nummerierung
            # st.markdown(f'<div class="section-header">📄 Text Passage {passage_index + 1}</div>', unsafe_allow_html=True)
            
            # Get all citation values for this passage
            citations = []
            if passage_index in st.session_state.annotations:
                citations = [
                    c for c in st.session_state.annotations[passage_index].get('Citation', []) 
                    if c and str(c).strip()
                ]
            
            # Show highlighted text passage
            highlighted_text_passage = highlight_citations(passage_text, citations)
            st.markdown(
                f'<div class="text-passage">{highlighted_text_passage}</div>',
                unsafe_allow_html=True
            )

            # Editable text area
            new_text = st.text_area(
                f"Edit Text Passage {passage_index + 1}",
                value=passage_text,
                height=100,
                key=f"text_passage_{passage_index}",
                label_visibility="collapsed"
            )
            
            # Update text in session state
            st.session_state.text_passages[passage_index] = new_text

            # Remove passage button (only if more than one passage exists)
            if len(st.session_state.text_passages) > 1:
                if st.button(f"🗑️ Remove Passage {passage_index + 1}", 
                           key=f"remove_passage_{passage_index}"):
                    remove_text_passage(passage_index)
                    st.rerun()
        
        with col2:
            # Annotations section
            # Entferne die Überschrift mit Nummerierung
            # st.markdown(f'<div class="section-header">📋 Annotations for Passage {passage_index + 1}</div>', unsafe_allow_html=True)
            
            # Ensure annotations exist for this passage
            if passage_index not in st.session_state.annotations:
                st.session_state.annotations[passage_index] = pd.DataFrame({
                    'Citation': [''],
                    'Metadata': [''],
                    'Annotation_Text': ['']
                })
            
            # Add the "Found in Text" column dynamically
            display_annotations = add_found_in_text_column(
                st.session_state.annotations[passage_index], 
                st.session_state.text_passages[passage_index]
            )
            
            # Editable annotations table with the new column
            edited_annotations = st.data_editor(
                display_annotations,
                use_container_width=True,
                num_rows="dynamic",
                key=f"annotations_{passage_index}",
                column_config={
                    "Found in Text": st.column_config.TextColumn(
                        "Found in Text",
                        help="Shows if citation is found in the text passage",
                        width="small",
                        disabled=True  # This column is read-only
                    ),
                    "Citation": st.column_config.TextColumn(
                        "Citation",
                        help="Source citation for verification",
                        width="medium"
                    ),
                    "Metadata": st.column_config.TextColumn(
                        "Metadata",
                        help="Source metadata (author, year, etc.)",
                        width="medium"
                    ),
                    "Annotation_Text": st.column_config.TextColumn(
                        "Annotation Text",
                        help="Explanations and comments",
                        width="large"
                    )
                },
                column_order=["Found in Text", "Citation", "Metadata", "Annotation_Text"]
            )
            
            # Update annotations in session state (remove the computed column)
            core_columns = ['Citation', 'Metadata', 'Annotation_Text']
            if all(col in edited_annotations.columns for col in core_columns):
                st.session_state.annotations[passage_index] = edited_annotations[core_columns].copy()
            
            # Annotation control buttons
            col_add, col_info = st.columns([1, 1])
            
            with col_add:
                if st.button(f"➕ Add Annotation", 
                           key=f"add_annotation_{passage_index}"):
                    add_annotation_row(passage_index)
                    st.rerun()
            
            with col_info:
                found_count = sum(1 for citation in st.session_state.annotations[passage_index]['Citation'] 
                                if check_citation_in_text(citation, st.session_state.text_passages[passage_index]))
                total_count = len(st.session_state.annotations[passage_index])
                st.caption(f"📊 {total_count} annotation(s) | ✅ {found_count} found in text")
        
        # Visual separator between passages
        if passage_index < len(st.session_state.text_passages) - 1:
            st.markdown("---")
    
    # Footer with statistics
    st.markdown("---")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📄 Text Passages", len(st.session_state.text_passages))
    
    with col2:
        total_annotations = sum(len(df) for df in st.session_state.annotations.values())
        st.metric("📋 Total Annotations", total_annotations)
    
    with col3:
        total_words = sum(len(text.split()) for text in st.session_state.text_passages)
        st.metric("📝 Total Words", total_words)
    
    with col4:
        # Calculate total found citations
        total_found = 0
        for passage_index, annotations_df in st.session_state.annotations.items():
            if passage_index < len(st.session_state.text_passages):
                passage_text = st.session_state.text_passages[passage_index]
                total_found += sum(1 for citation in annotations_df['Citation'] 
                                 if check_citation_in_text(citation, passage_text))
        st.metric("✅ Citations Found", total_found)

if __name__ == "__main__":
    main()
